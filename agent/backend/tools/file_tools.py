"""
File Processing Tools for Backend Agent
Handles file uploads, processing, format conversion, and data extraction.
"""

import os
import logging
import asyncio
import shutil
import mimetypes
import hashlib
import zipfile
import csv
import json
import xml.etree.ElementTree as ET
from typing import Dict, Any, List, Optional, Union, BinaryIO
from datetime import datetime
from pathlib import Path
from io import BytesIO, StringIO

import pandas as pd
import openpyxl
from PIL import Image, ImageOps
import PyPDF2
import docx
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

class FileInfo(BaseModel):
    """Model for file information"""
    filename: str = Field(..., description="Original filename")
    file_path: str = Field(..., description="File storage path")
    file_size: int = Field(..., description="File size in bytes")
    mime_type: str = Field(..., description="MIME type")
    checksum: str = Field(..., description="File checksum")
    uploaded_at: datetime = Field(default_factory=datetime.utcnow)

class ProcessingResult(BaseModel):
    """Model for file processing results"""
    success: bool = Field(..., description="Whether processing was successful")
    file_info: Optional[FileInfo] = None
    data: Any = Field(None, description="Processed data")
    message: str = Field("", description="Processing message")
    processing_time: Optional[float] = Field(None, description="Processing time in seconds")

class FileProcessingTool:
    """Advanced file processing tool with support for multiple formats"""
    
    def __init__(self):
        self.upload_directory = Path(os.getenv('UPLOAD_DIRECTORY', './uploads'))
        self.max_file_size = int(os.getenv('MAX_FILE_SIZE', 10485760))  # 10MB
        self.allowed_extensions = set(os.getenv('ALLOWED_EXTENSIONS', 'txt,pdf,docx,xlsx,csv,json,xml,jpg,png,gif').split(','))
        
        # Create upload directory if it doesn't exist
        self.upload_directory.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"File processing tool initialized - Upload dir: {self.upload_directory}")

    def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate MD5 checksum of file"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def _validate_file(self, filename: str, file_size: int) -> Optional[str]:
        """Validate file before processing"""
        # Check file extension
        file_extension = Path(filename).suffix.lower().lstrip('.')
        if file_extension not in self.allowed_extensions:
            return f"File type '{file_extension}' not allowed. Allowed types: {', '.join(self.allowed_extensions)}"
        
        # Check file size
        if file_size > self.max_file_size:
            return f"File size ({file_size} bytes) exceeds maximum allowed size ({self.max_file_size} bytes)"
        
        return None

    async def upload_file(self, filename: str, file_content: bytes) -> ProcessingResult:
        """Upload and store file"""
        start_time = datetime.utcnow()
        
        try:
            # Validate file
            validation_error = self._validate_file(filename, len(file_content))
            if validation_error:
                return ProcessingResult(
                    success=False,
                    message=validation_error
                )
            
            # Generate unique filename to prevent conflicts
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            file_stem = Path(filename).stem
            file_extension = Path(filename).suffix
            unique_filename = f"{file_stem}_{timestamp}{file_extension}"
            
            # Save file
            file_path = self.upload_directory / unique_filename
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # Get file information
            mime_type, _ = mimetypes.guess_type(filename)
            checksum = self._calculate_checksum(file_path)
            
            file_info = FileInfo(
                filename=filename,
                file_path=str(file_path),
                file_size=len(file_content),
                mime_type=mime_type or "application/octet-stream",
                checksum=checksum
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"File uploaded successfully: {unique_filename} ({len(file_content)} bytes)")
            
            return ProcessingResult(
                success=True,
                file_info=file_info,
                message=f"File uploaded successfully as {unique_filename}",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"File upload failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"File upload failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_csv_file(self, file_path: str, delimiter: str = ',') -> ProcessingResult:
        """Process CSV file and return data"""
        start_time = datetime.utcnow()
        
        try:
            # Read CSV file
            df = pd.read_csv(file_path, delimiter=delimiter)
            
            # Convert to dictionary format
            data = {
                "columns": df.columns.tolist(),
                "rows": df.to_dict('records'),
                "row_count": len(df),
                "column_count": len(df.columns),
                "summary": df.describe().to_dict() if df.select_dtypes(include=['number']).empty == False else {}
            }
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"CSV processed successfully: {len(df)} rows, {len(df.columns)} columns")
            
            return ProcessingResult(
                success=True,
                data=data,
                message=f"CSV processed successfully: {len(df)} rows",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"CSV processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"CSV processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_excel_file(self, file_path: str, sheet_name: str = None) -> ProcessingResult:
        """Process Excel file and return data"""
        start_time = datetime.utcnow()
        
        try:
            # Read Excel file
            if sheet_name:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheets_data = {sheet_name: df}
            else:
                # Read all sheets
                excel_file = pd.ExcelFile(file_path)
                sheets_data = {}
                for sheet in excel_file.sheet_names:
                    sheets_data[sheet] = pd.read_excel(file_path, sheet_name=sheet)
            
            # Process each sheet
            result_data = {}
            total_rows = 0
            
            for sheet, df in sheets_data.items():
                result_data[sheet] = {
                    "columns": df.columns.tolist(),
                    "rows": df.to_dict('records'),
                    "row_count": len(df),
                    "column_count": len(df.columns)
                }
                total_rows += len(df)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"Excel processed successfully: {len(sheets_data)} sheets, {total_rows} total rows")
            
            return ProcessingResult(
                success=True,
                data=result_data,
                message=f"Excel processed successfully: {len(sheets_data)} sheets",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"Excel processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"Excel processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_pdf_file(self, file_path: str) -> ProcessingResult:
        """Extract text from PDF file"""
        start_time = datetime.utcnow()
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        text_content.append({
                            "page": page_num + 1,
                            "text": text.strip()
                        })
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        text_content.append({
                            "page": page_num + 1,
                            "text": "",
                            "error": str(e)
                        })
            
            # Combine all text
            full_text = "\n".join([page["text"] for page in text_content if page["text"]])
            
            data = {
                "pages": text_content,
                "page_count": len(text_content),
                "full_text": full_text,
                "word_count": len(full_text.split()),
                "character_count": len(full_text)
            }
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"PDF processed successfully: {len(text_content)} pages")
            
            return ProcessingResult(
                success=True,
                data=data,
                message=f"PDF processed successfully: {len(text_content)} pages",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"PDF processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"PDF processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_docx_file(self, file_path: str) -> ProcessingResult:
        """Extract text from DOCX file"""
        start_time = datetime.utcnow()
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            full_text = "\n".join(paragraphs)
            
            data = {
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "full_text": full_text,
                "word_count": len(full_text.split()),
                "character_count": len(full_text)
            }
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"DOCX processed successfully: {len(paragraphs)} paragraphs, {len(tables)} tables")
            
            return ProcessingResult(
                success=True,
                data=data,
                message=f"DOCX processed successfully",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"DOCX processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"DOCX processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_image_file(self, file_path: str, operations: List[str] = None) -> ProcessingResult:
        """Process image file with optional operations"""
        start_time = datetime.utcnow()
        
        try:
            with Image.open(file_path) as img:
                # Get image info
                image_info = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height
                }
                
                processed_images = {"original": image_info}
                
                # Apply operations if specified
                if operations:
                    for operation in operations:
                        try:
                            if operation == "thumbnail":
                                # Create thumbnail
                                thumb = img.copy()
                                thumb.thumbnail((200, 200))
                                thumb_path = file_path.replace('.', '_thumb.')
                                thumb.save(thumb_path)
                                processed_images["thumbnail"] = {
                                    "path": thumb_path,
                                    "size": thumb.size
                                }
                            
                            elif operation == "grayscale":
                                # Convert to grayscale
                                gray = ImageOps.grayscale(img)
                                gray_path = file_path.replace('.', '_gray.')
                                gray.save(gray_path)
                                processed_images["grayscale"] = {
                                    "path": gray_path,
                                    "size": gray.size
                                }
                            
                            elif operation == "resize_50":
                                # Resize to 50%
                                new_size = (img.width // 2, img.height // 2)
                                resized = img.resize(new_size)
                                resized_path = file_path.replace('.', '_50.')
                                resized.save(resized_path)
                                processed_images["resized_50"] = {
                                    "path": resized_path,
                                    "size": resized.size
                                }
                                
                        except Exception as e:
                            logger.warning(f"Image operation '{operation}' failed: {str(e)}")
                
                processing_time = (datetime.utcnow() - start_time).total_seconds()
                
                logger.info(f"Image processed successfully: {img.width}x{img.height} {img.format}")
                
                return ProcessingResult(
                    success=True,
                    data=processed_images,
                    message=f"Image processed successfully",
                    processing_time=processing_time
                )
                
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"Image processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"Image processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def process_json_file(self, file_path: str) -> ProcessingResult:
        """Process JSON file and return structured data"""
        start_time = datetime.utcnow()
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            
            # Analyze JSON structure
            def analyze_structure(obj, path="root"):
                if isinstance(obj, dict):
                    return {
                        "type": "object",
                        "keys": list(obj.keys()),
                        "key_count": len(obj),
                        "children": {k: analyze_structure(v, f"{path}.{k}") for k, v in obj.items()}
                    }
                elif isinstance(obj, list):
                    return {
                        "type": "array",
                        "length": len(obj),
                        "item_types": list(set(type(item).__name__ for item in obj)),
                        "sample": obj[:3] if obj else []
                    }
                else:
                    return {
                        "type": type(obj).__name__,
                        "value": obj if not isinstance(obj, (str)) or len(str(obj)) < 100 else str(obj)[:100] + "..."
                    }
            
            structure = analyze_structure(json_data)
            
            data = {
                "content": json_data,
                "structure": structure,
                "size_bytes": os.path.getsize(file_path)
            }
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"JSON processed successfully")
            
            return ProcessingResult(
                success=True,
                data=data,
                message="JSON processed successfully",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"JSON processing failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"JSON processing failed: {str(e)}",
                processing_time=processing_time
            )

    async def generate_csv_export(self, data: List[Dict[str, Any]], filename: str = None) -> ProcessingResult:
        """Generate CSV file from data"""
        start_time = datetime.utcnow()
        
        try:
            if not data:
                return ProcessingResult(
                    success=False,
                    message="No data provided for CSV export"
                )
            
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                filename = f"export_{timestamp}.csv"
            
            file_path = self.upload_directory / filename
            
            # Write CSV
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                if data:
                    fieldnames = data[0].keys()
                    writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(data)
            
            file_size = os.path.getsize(file_path)
            checksum = self._calculate_checksum(file_path)
            
            file_info = FileInfo(
                filename=filename,
                file_path=str(file_path),
                file_size=file_size,
                mime_type="text/csv",
                checksum=checksum
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"CSV export generated: {filename} ({len(data)} rows)")
            
            return ProcessingResult(
                success=True,
                file_info=file_info,
                data={"row_count": len(data), "filename": filename},
                message=f"CSV export generated successfully: {len(data)} rows",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"CSV export failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"CSV export failed: {str(e)}",
                processing_time=processing_time
            )

    async def create_zip_archive(self, file_paths: List[str], archive_name: str = None) -> ProcessingResult:
        """Create ZIP archive from multiple files"""
        start_time = datetime.utcnow()
        
        try:
            if not archive_name:
                timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                archive_name = f"archive_{timestamp}.zip"
            
            archive_path = self.upload_directory / archive_name
            
            with zipfile.ZipFile(archive_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file_path in file_paths:
                    if os.path.exists(file_path):
                        # Add file to archive with just the filename
                        arcname = os.path.basename(file_path)
                        zipf.write(file_path, arcname)
                    else:
                        logger.warning(f"File not found for archive: {file_path}")
            
            file_size = os.path.getsize(archive_path)
            checksum = self._calculate_checksum(archive_path)
            
            file_info = FileInfo(
                filename=archive_name,
                file_path=str(archive_path),
                file_size=file_size,
                mime_type="application/zip",
                checksum=checksum
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            logger.info(f"ZIP archive created: {archive_name} ({len(file_paths)} files)")
            
            return ProcessingResult(
                success=True,
                file_info=file_info,
                data={"file_count": len(file_paths), "archive_name": archive_name},
                message=f"ZIP archive created successfully",
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            logger.error(f"ZIP archive creation failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"ZIP archive creation failed: {str(e)}",
                processing_time=processing_time
            )

    async def delete_file(self, file_path: str) -> ProcessingResult:
        """Delete file from storage"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
                return ProcessingResult(
                    success=True,
                    message=f"File deleted successfully"
                )
            else:
                return ProcessingResult(
                    success=False,
                    message="File not found"
                )
                
        except Exception as e:
            logger.error(f"File deletion failed: {str(e)}")
            return ProcessingResult(
                success=False,
                message=f"File deletion failed: {str(e)}"
            )

    def health_check(self) -> Dict[str, Any]:
        """Perform health check on file processing system"""
        try:
            # Check upload directory
            upload_dir_exists = self.upload_directory.exists()
            upload_dir_writable = os.access(self.upload_directory, os.W_OK) if upload_dir_exists else False
            
            # Count files in upload directory
            file_count = len(list(self.upload_directory.glob('*'))) if upload_dir_exists else 0
            
            return {
                "status": "healthy" if upload_dir_exists and upload_dir_writable else "unhealthy",
                "upload_directory": str(self.upload_directory),
                "upload_dir_exists": upload_dir_exists,
                "upload_dir_writable": upload_dir_writable,
                "file_count": file_count,
                "max_file_size": self.max_file_size,
                "allowed_extensions": list(self.allowed_extensions)
            }
        except Exception as e:
            return {
                "status": "unhealthy",
                "error": str(e)
            }

# Factory function to create file processing tools
def create_file_tools() -> List:
    """Create and return file processing tools for the agent"""
    file_tool = FileProcessingTool()
    
    # Define individual tool functions that the agent can use
    async def upload_and_process_file(filename: str, file_content: bytes, process_type: str = "auto") -> dict:
        """Upload file and process based on type"""
        # First upload the file
        upload_result = await file_tool.upload_file(filename, file_content)
        
        if not upload_result.success:
            return upload_result.model_dump()
        
        file_path = upload_result.file_info.file_path
        
        # Process based on file type or explicit process_type
        file_extension = Path(filename).suffix.lower().lstrip('.')
        
        if process_type == "auto":
            if file_extension == 'csv':
                process_result = await file_tool.process_csv_file(file_path)
            elif file_extension in ['xlsx', 'xls']:
                process_result = await file_tool.process_excel_file(file_path)
            elif file_extension == 'pdf':
                process_result = await file_tool.process_pdf_file(file_path)
            elif file_extension == 'docx':
                process_result = await file_tool.process_docx_file(file_path)
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
                process_result = await file_tool.process_image_file(file_path)
            elif file_extension == 'json':
                process_result = await file_tool.process_json_file(file_path)
            else:
                return upload_result.model_dump()
        else:
            # Process based on explicit type
            if process_type == "csv":
                process_result = await file_tool.process_csv_file(file_path)
            elif process_type == "excel":
                process_result = await file_tool.process_excel_file(file_path)
            elif process_type == "pdf":
                process_result = await file_tool.process_pdf_file(file_path)
            elif process_type == "image":
                process_result = await file_tool.process_image_file(file_path, ["thumbnail", "grayscale"])
            else:
                return upload_result.model_dump()
        
        # Combine upload and processing results
        combined_result = upload_result.model_dump()
        combined_result["processing_result"] = process_result.model_dump()
        
        return combined_result
    
    async def export_data_to_csv(data: list, filename: str = None) -> dict:
        """Export data to CSV file"""
        result = await file_tool.generate_csv_export(data, filename)
        return result.model_dump()
    
    async def create_file_archive(file_paths: list, archive_name: str = None) -> dict:
        """Create ZIP archive from files"""
        result = await file_tool.create_zip_archive(file_paths, archive_name)
        return result.model_dump()
    
    async def remove_file(file_path: str) -> dict:
        """Delete file from storage"""
        result = await file_tool.delete_file(file_path)
        return result.model_dump()
    
    return [
        upload_and_process_file,
        export_data_to_csv,
        create_file_archive,
        remove_file
    ]

if __name__ == "__main__":
    # Test the file processing tool
    async def test_file_tool():
        file_tool = FileProcessingTool()
        
        # Test health check
        health = file_tool.health_check()
        print(f"Health check: {health}")
        
        # Test CSV generation
        test_data = [
            {"name": "John", "age": 30, "city": "New York"},
            {"name": "Jane", "age": 25, "city": "San Francisco"},
            {"name": "Bob", "age": 35, "city": "Chicago"}
        ]
        
        result = await file_tool.generate_csv_export(test_data, "test_export.csv")
        print(f"CSV export result: {result}")
    
    asyncio.run(test_file_tool())
